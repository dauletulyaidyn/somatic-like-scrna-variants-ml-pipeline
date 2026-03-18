#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]  # .../Final version
TABLES = ROOT / "results" / "tables"
FIGURES = ROOT / "results" / "figures"
OUT_DOCX = ROOT / "manuscript" / "manuscript_sadhana_ml_it.docx"


def set_document_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)


def add_par(doc: Document, text: str, align: WD_ALIGN_PARAGRAPH | None = None, bold: bool = False) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold


def add_table_from_df(doc: Document, title: str, df: pd.DataFrame) -> None:
    add_par(doc, title, bold=True)
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, c in enumerate(df.columns):
        t.rows[0].cells[i].text = str(c)

    for _, row in df.iterrows():
        cells = t.add_row().cells
        for i, c in enumerate(df.columns):
            v = row[c]
            if isinstance(v, float):
                if abs(v) < 0.001 and v != 0:
                    cells[i].text = f"{v:.2e}"
                else:
                    cells[i].text = f"{v:.6f}".rstrip("0").rstrip(".")
            else:
                cells[i].text = str(v)

    doc.add_paragraph("")


def add_figure(doc: Document, idx: int, filename: str, caption: str) -> None:
    path = FIGURES / filename
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(6.5))
    add_par(doc, f"Figure {idx}. {caption}")
    doc.add_paragraph("")


def load_core_tables() -> dict[str, pd.DataFrame]:
    tables = {
        "table1": pd.read_csv(TABLES / "table1_sample_summary.csv"),
        "table2": pd.read_csv(TABLES / "table2_variant_counts_summary.csv"),
        "table3": pd.read_csv(TABLES / "table3_ml_summary.csv"),
        "leakage": pd.read_csv(TABLES / "table_leakage_sensitivity.csv"),
        "metrics_repeated": pd.read_csv(TABLES / "supervised_metrics_repeated.csv"),
    }
    cmp_path = TABLES / "table_bcftools_vs_gatk_summary.csv"
    if cmp_path.exists():
        tables["caller_comparison"] = pd.read_csv(cmp_path)
    return tables


def extract_metrics(core: dict[str, pd.DataFrame]) -> dict[str, float | int]:
    t1 = core["table1"]
    t2 = core["table2"]
    t3 = core["table3"]
    leak = core["leakage"]

    n_total = int(t1["n_samples"].sum())
    n_uwe = int(t1.loc[t1["condition"] == "unwounded_skin", "n_samples"].iloc[0])
    n_we = int(t1.loc[t1["condition"] == "wound_edge", "n_samples"].iloc[0])

    t2f = t2.loc[t2["set"] == "filtered"].iloc[0]
    t2c = t2.loc[t2["set"] == "cohort_filtered"].iloc[0]
    t3f = t3.loc[t3["feature_set"] == "filtered"].iloc[0]
    t3c = t3.loc[t3["feature_set"] == "cohort_filtered"].iloc[0]

    leak_run = leak.loc[leak["evaluation"] == "run_level_repeated_stratified_5x30"].iloc[0]
    leak_donor = leak.loc[leak["evaluation"] == "donor_group_kfold_7"].iloc[0]

    return {
        "n_total": n_total,
        "n_uwe": n_uwe,
        "n_we": n_we,
        "records_filtered_median": int(t2f["records_median"]),
        "records_filtered_min": int(t2f["records_min"]),
        "records_filtered_max": int(t2f["records_max"]),
        "records_cohort_median": int(t2c["records_median"]),
        "records_cohort_min": int(t2c["records_min"]),
        "records_cohort_max": int(t2c["records_max"]),
        "features_filtered": int(t3f["features"]),
        "features_cohort": int(t3c["features"]),
        "bacc_filtered": float(t3f["best_balanced_accuracy_mean"]),
        "bacc_cohort": float(t3c["best_balanced_accuracy_mean"]),
        "p_filtered": float(t3f["permutation_pvalue"]),
        "p_cohort": float(t3c["permutation_pvalue"]),
        "repeats": int(t3f["repeats"]),
        "permutations": int(t3f["permutations"]),
        "n_unique_donors": int(leak_run["n_unique_donors"]),
        "run_level_bacc": float(leak_run["balanced_accuracy_mean"]),
        "donor_level_bacc": float(leak_donor["balanced_accuracy_mean"]),
    }


def build_enrichment_table() -> pd.DataFrame:
    rows = []
    files = [
        ("top_mutated_genes", "enrichment_results_top_mutated_genes.csv"),
        ("wound_edge_higher_mutation_burden", "enrichment_results_wound_edge_higher_mutation_burden.csv"),
        ("unwounded_skin_higher_mutation_burden", "enrichment_results_unwounded_skin_higher_mutation_burden.csv"),
    ]
    for label, fn in files:
        df = pd.read_csv(TABLES / fn).sort_values("p_value").head(4)
        df = df[["name", "source", "native", "p_value"]].copy()
        df.insert(0, "gene_set", label)
        rows.append(df)
    out = pd.concat(rows, ignore_index=True)
    return out


def to_display_samples(df: pd.DataFrame) -> pd.DataFrame:
    disp = df.copy()
    disp["condition"] = disp["condition"].map(
        {"unwounded_skin": "Unwounded skin", "wound_edge": "Wound edge"}
    )
    disp = disp.rename(columns={"condition": "Condition", "n_samples": "n"})
    return disp


def to_display_variants(df: pd.DataFrame) -> pd.DataFrame:
    disp = df.copy()
    disp["set"] = disp["set"].map({"filtered": "Filtered", "cohort_filtered": "Cohort-filtered"})
    disp = disp.rename(
        columns={
            "set": "Feature set",
            "n_samples": "n",
            "records_median": "Records (median)",
            "records_min": "Records (min)",
            "records_max": "Records (max)",
            "snps_median": "SNPs (median)",
            "indels_median": "Indels (median)",
        }
    )
    return disp


def to_display_ml(df: pd.DataFrame) -> pd.DataFrame:
    disp = df.copy()
    disp["feature_set"] = disp["feature_set"].map(
        {"filtered": "Filtered", "cohort_filtered": "Cohort-filtered"}
    )
    disp = disp.rename(
        columns={
            "feature_set": "Feature set",
            "samples": "Samples",
            "features": "Features",
            "best_balanced_accuracy_mean": "Best balanced accuracy",
            "permutation_pvalue": "Permutation p-value",
            "best_model_repeated": "Best repeated-CV model",
            "repeats": "Repeats",
            "permutations": "Permutations",
        }
    )
    return disp


def to_display_leakage(df: pd.DataFrame) -> pd.DataFrame:
    disp = df.copy()
    disp["evaluation"] = disp["evaluation"].map(
        {
            "run_level_repeated_stratified_5x30": "Run-level repeated stratified CV (5x30)",
            "donor_group_kfold_7": "Donor-grouped KFold (7 folds)",
        }
    )
    disp = disp.rename(
        columns={
            "evaluation": "Evaluation protocol",
            "n_samples": "n_samples",
            "n_unique_donors": "n_donors",
            "accuracy_mean": "Accuracy mean",
            "balanced_accuracy_mean": "Balanced accuracy mean",
        }
    )
    return disp


def to_display_models(df: pd.DataFrame) -> pd.DataFrame:
    disp = (
        df.sort_values("balanced_accuracy_mean", ascending=False)
        .head(6)
        .loc[:, ["model", "balanced_accuracy_mean", "balanced_accuracy_std", "roc_auc_mean"]]
        .copy()
    )
    disp = disp.rename(
        columns={
            "model": "Model",
            "balanced_accuracy_mean": "Balanced accuracy mean",
            "balanced_accuracy_std": "Balanced accuracy std",
            "roc_auc_mean": "ROC AUC mean",
        }
    )
    return disp


def main() -> int:
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)

    core = load_core_tables()
    m = extract_metrics(core)
    enrichment = build_enrichment_table()
    top_models = to_display_models(core["metrics_repeated"])

    doc = Document()
    set_document_style(doc)

    title = (
        "A reproducible ML pipeline for scRNA-seq variant-derived classification: "
        "wound edge versus unwounded skin"
    )
    running_title = "ML Pipeline for scRNA Variant Classification"

    add_par(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_par(doc, "Kunikeyev Aidyn1,*", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_par(doc, "1 Affiliation: not provided in workspace files", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_par(doc, "*For correspondence: email not provided in workspace files", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_par(doc, "ORCID IDs (recommended by journal): not provided", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_par(doc, f"Running title: {running_title}", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    doc.add_heading("Abstract", level=1)
    add_par(
        doc,
        "This manuscript presents a reproducible machine-learning/data-science workflow that transforms "
        "single-cell RNA-seq expressed-variant calls into sample-level features for binary classification. "
        "The analysis used PRJNA736095 with "
        f"{m['n_total']} runs ({m['n_uwe']} unwounded skin, {m['n_we']} wound edge). "
        "The pipeline combines STARsolo alignment, GATK RNA variant calling, cohort-frequency filtering, "
        "variant-to-gene feature engineering, repeated cross-validation, and permutation testing. "
        f"Filtered VCFs had median {m['records_filtered_median']:,} records per sample "
        f"(range {m['records_filtered_min']:,}-{m['records_filtered_max']:,}); cohort-filtered VCFs had median "
        f"{m['records_cohort_median']:,} (range {m['records_cohort_min']:,}-{m['records_cohort_max']:,}). "
        "Best repeated-CV performance (balanced accuracy) was "
        f"{m['bacc_filtered']:.3f} for filtered features and {m['bacc_cohort']:.3f} for cohort-filtered features, "
        f"with permutation p-values {m['p_filtered']:.5f} and {m['p_cohort']:.5f}, respectively "
        f"({m['repeats']} repeats; {m['permutations']} permutations). "
        "A leakage-sensitivity analysis using donor-grouped folds showed a marked drop in balanced accuracy "
        f"from {m['run_level_bacc']:.3f} (run-level CV) to {m['donor_level_bacc']:.3f} (donor-grouped CV), "
        "highlighting the importance of split design. "
        "Bioinformatics interpretation is intentionally minimal and restricted to compact enrichment summaries."
    )
    add_par(
        doc,
        "MS received [to be filled]; revised [to be filled]; accepted [to be filled]"
    )

    doc.add_heading("Keywords", level=1)
    add_par(
        doc,
        "Keywords. scRNA-seq, variant calling, feature engineering, machine learning, cross-validation, reproducibility"
    )

    doc.add_heading("List of Symbols and Abbreviations", level=1)
    add_par(doc, "n: number of samples")
    add_par(doc, "p: number of features (genes)")
    add_par(doc, "X: sample-by-feature matrix (gene-level variant burden)")
    add_par(doc, "y: binary class label (unwounded_skin, wound_edge)")
    add_par(doc, "CV: cross-validation")
    add_par(doc, "DP: read depth; AD: alternate allele depth; VAF: variant allele fraction")

    doc.add_heading("1 Introduction", level=1)
    add_par(
        doc,
        "From an IT/ML perspective, this work addresses a practical problem: how to convert sparse, noisy, "
        "high-dimensional sequencing outputs into reproducible classification features while controlling for "
        "small-sample overfitting. The main contribution is not a biological claim, but a fully auditable "
        "data-processing and model-validation pipeline. The workflow uses standard tools for alignment and "
        "variant calling [1-4], then applies strict evaluation protocols from machine learning [5,6]."
    )

    doc.add_heading("2 Materials and Methods", level=1)

    doc.add_heading("2.1 Dataset and Labels", level=2)
    add_par(
        doc,
        f"The dataset is BioProject PRJNA736095 with {m['n_total']} sequencing runs mapped to two classes: "
        f"{m['n_uwe']} unwounded skin and {m['n_we']} wound edge. A donor-linked metadata table identifies "
        f"{m['n_unique_donors']} unique donors (two runs per donor in this subset)."
    )

    doc.add_heading("2.2 Data Processing and Variant Calling", level=2)
    add_par(
        doc,
        "FASTQ reads were aligned with STARsolo [1] to produce coordinate-sorted BAM files. "
        "Per-sample variant calling used the GATK RNA workflow centered on HaplotypeCaller [3] on aligned reads. "
        "A quality filter retained variants satisfying DP>=10, QUAL>=30, AD_alt>=3, and VAF>=0.10. "
        "A cohort-frequency filter then excluded loci observed in >=4 samples (cohort-common positions) "
        "to define a somatic-like candidate set. This step is a computational proxy and not a matched-normal "
        "somatic confirmation."
    )

    doc.add_heading("2.3 Feature Engineering", level=2)
    add_par(
        doc,
        "Variants were intersected with exon coordinates using BEDTools [4] to produce variant-to-gene mappings. "
        "For each sample and gene, the feature value is the count of unique variant keys "
        "(chrom:pos:ref>alt), yielding a sparse gene-level burden matrix X."
    )

    doc.add_heading("2.4 Supervised Modeling and Validation", level=2)
    add_par(
        doc,
        "Models were implemented with scikit-learn [5]: logistic regression (L1/L2), linear SVC, "
        "random forest, logistic regression with PCA, and logistic regression with univariate feature selection. "
        "The main estimate used RepeatedStratifiedKFold with 5 folds and 30 repeats, scoring balanced accuracy. "
        "A permutation test with 200 permutations followed the framework in [6] to assess whether observed "
        "performance exceeded label-randomized baselines."
    )

    doc.add_heading("2.5 Leakage-Sensitivity Analysis", level=2)
    add_par(
        doc,
        "To test split dependence, we compared run-level repeated stratified CV versus donor-grouped KFold "
        "(7 folds). This analysis quantifies how non-independent runs from the same donor can inflate run-level "
        "estimates if grouping is ignored."
    )

    doc.add_heading("2.6 Minimal Functional Context", level=2)
    add_par(
        doc,
        "Functional enrichment (g:Profiler [11]) was applied only as a compact context layer. "
        "A minimal single-cell context layer (Scanpy/Leiden/UMAP with cellsnp-lite) is retained for reporting "
        "consistency [7-10]. Interpretation is intentionally limited in this IT-focused manuscript and can be "
        "expanded by domain specialists."
    )

    doc.add_heading("3 Results", level=1)

    doc.add_heading("3.1 Cohort and Variant Yield", level=2)
    add_par(
        doc,
        f"Across {m['n_total']} runs, the median number of filtered variant records per sample was "
        f"{m['records_filtered_median']:,}. After cohort-frequency filtering, the median was "
        f"{m['records_cohort_median']:,}, indicating substantial reduction of recurrent positions."
    )

    add_table_from_df(doc, "Table 1. Sample counts by class.", to_display_samples(core["table1"]))
    add_table_from_df(
        doc,
        "Table 2. Variant count summary before and after cohort-frequency filtering.",
        to_display_variants(core["table2"]),
    )
    if "caller_comparison" in core:
        add_table_from_df(doc, "Table S1. Historical baseline vs active GATK comparison.", core["caller_comparison"])

    doc.add_heading("3.2 Supervised Classification Performance", level=2)
    add_par(
        doc,
        f"Gene-level feature dimensions were {m['features_filtered']:,} (filtered) and "
        f"{m['features_cohort']:,} (cohort-filtered). Best repeated-CV balanced accuracy reached "
        f"{m['bacc_filtered']:.3f} and {m['bacc_cohort']:.3f}, with permutation p-values "
        f"{m['p_filtered']:.5f} and {m['p_cohort']:.5f}, respectively."
    )

    add_table_from_df(doc, "Table 3. Main ML summary.", to_display_ml(core["table3"]))
    add_table_from_df(
        doc,
        "Table 4. Top models in repeated CV (balanced accuracy ranking).",
        top_models,
    )

    doc.add_heading("3.3 Leakage-Sensitivity Outcome", level=2)
    add_par(
        doc,
        f"Run-level repeated stratified CV yielded balanced accuracy {m['run_level_bacc']:.3f}. "
        f"When folds were constrained by donor, balanced accuracy decreased to {m['donor_level_bacc']:.3f}. "
        "This confirms that split strategy is a first-order factor for realistic performance estimation."
    )

    add_table_from_df(
        doc,
        "Table 5. Leakage-sensitivity comparison (run-level vs donor-grouped splits).",
        to_display_leakage(core["leakage"]),
    )

    doc.add_heading("3.4 Minimal Functional Interpretation", level=2)
    add_par(
        doc,
        "Only high-level enrichment patterns are reported: top-mutated genes were enriched for broad cellular "
        "component and binding categories; wound-edge higher burden sets were enriched for antigen/MHC-related "
        "terms; unwounded-skin higher burden sets were enriched for scavenger receptor and FCGR-associated pathways."
    )
    add_table_from_df(
        doc,
        "Table 6. Top enrichment terms per gene-set category (lowest p-values).",
        enrichment.rename(
            columns={
                "gene_set": "Gene set",
                "name": "Term",
                "source": "Source",
                "native": "Native ID",
                "p_value": "p-value",
            }
        ),
    )

    doc.add_heading("4 Discussion", level=1)
    add_par(
        doc,
        "The primary outcome is an engineering result: a reproducible path from sequencing artifacts to "
        "validated ML outputs, including explicit checks against optimistic split designs. In small cohorts, "
        "reporting only high run-level metrics is unsafe; repeated CV, permutation testing, and grouped-split "
        "sensitivity should be standard practice. This framework can be reused for other two-class tasks with "
        "similar high-dimensional sparse features."
    )

    doc.add_heading("5 Conclusions", level=1)
    add_par(
        doc,
        "A complete IT-oriented pipeline was built and documented for variant-derived classification from scRNA-seq. "
        "The workflow demonstrates strong run-level discrimination but also quantifies donor-level performance drop, "
        "which is critical for honest generalization claims. Bioinformatics interpretation is kept minimal by design."
    )

    doc.add_heading("Appendix A", level=1)
    add_par(
        doc,
        "Reproducibility assets include configuration files, generated tables, and figures under "
        "U:/PAD/Final version/results/ and the manuscript generator script under "
        "U:/PAD/Final version/repo/scripts/."
    )

    doc.add_heading("Acknowledgement", level=1)
    add_par(
        doc,
        "Public data providers (SRA/GEO maintainers and original PRJNA736095 contributors) are acknowledged. "
        "Funding details were not provided in workspace files."
    )

    doc.add_heading("Data and Code Availability", level=1)
    add_par(doc, "Data accession: PRJNA736095.")
    add_par(doc, "Results used in this manuscript: U:/PAD/Final version/results/.")
    add_par(
        doc,
        "Pipeline repository referenced in project materials: "
        "https://github.com/dauletulyaidyn/somatic-like-scrna-variants-ml-pipeline"
    )

    doc.add_heading("Figures", level=1)
    figure_specs = [
        ("fig1_pipeline_flow.png", "End-to-end computational pipeline."),
        ("fig2_variant_counts_filtered.png", "Per-sample filtered variant counts."),
        ("fig3_variant_counts_cohort_filtered.png", "Per-sample cohort-filtered variant counts."),
        ("fig6_ml_repeatedcv_baseline.png", "Repeated-CV model comparison (filtered features)."),
        ("fig7_ml_repeatedcv_cohort.png", "Repeated-CV model comparison (cohort-filtered features)."),
        ("fig7_pca_filtered_class_donor.png", "PCA view with class and donor overlays (filtered features)."),
        ("fig8_pca_cohort_class_donor.png", "PCA view with class and donor overlays (cohort-filtered features)."),
        ("fig10_enrichment_top_terms.png", "Top enrichment terms (compact summary)."),
    ]

    for idx, (fname, cap) in enumerate(figure_specs, start=1):
        add_figure(doc, idx, fname, cap)

    doc.add_heading("Figure Captions", level=1)
    for idx, (_, cap) in enumerate(figure_specs, start=1):
        add_par(doc, f"Figure {idx}. {cap}")

    doc.add_heading("References", level=1)
    refs = [
        "[1] Dobin A, Davis CA, Schlesinger F, et al. 2013 STAR: ultrafast universal RNA-seq aligner. Bioinformatics 29(1):15-21.",
        "[2] Li H, Handsaker B, Wysoker A, et al. 2009 The Sequence Alignment/Map format and SAMtools. Bioinformatics 25(16):2078-2079.",
        "[3] Danecek P, Bonfield JK, Liddle J, et al. 2021 Twelve years of SAMtools and BCFtools. GigaScience 10(2):giab008.",
        "[4] Quinlan AR and Hall IM. 2010 BEDTools: a flexible suite of utilities for comparing genomic features. Bioinformatics 26(6):841-842.",
        "[5] Pedregosa F, Varoquaux G, Gramfort A, et al. 2011 Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research 12:2825-2830.",
        "[6] Ojala M and Garriga GC. 2010 Permutation Tests for Studying Classifier Performance. Journal of Machine Learning Research 11:1833-1863.",
        "[7] Wolf FA, Angerer P and Theis FJ. 2018 SCANPY: large-scale single-cell gene expression data analysis. Genome Biology 19:15.",
        "[8] Traag VA, Waltman L and van Eck NJ. 2019 From Louvain to Leiden: guaranteeing well-connected communities. Scientific Reports 9(1):5233.",
        "[9] McInnes L, Healy J and Melville J. 2018 UMAP: Uniform Manifold Approximation and Projection for Dimension Reduction. arXiv:1802.03426.",
        "[10] Huang X and Huang Y. 2021 Cellsnp-lite: an efficient tool for genotyping single cells. Bioinformatics 37(23):4569-4571.",
        "[11] Raudvere U, Kolberg L, Kuzmin I, et al. 2019 g:Profiler: a web server for functional enrichment analysis and conversions of gene lists (2019 update). Nucleic Acids Research 47(W1):W191-W198.",
    ]
    for r in refs:
        add_par(doc, r)

    doc.save(str(OUT_DOCX))
    print(f"Wrote: {OUT_DOCX}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
