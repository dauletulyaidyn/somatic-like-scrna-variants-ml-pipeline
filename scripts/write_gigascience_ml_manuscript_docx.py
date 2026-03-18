#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path
from shutil import copy2

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[2]  # .../Final version
TABLES = ROOT / "results" / "tables"
FIGURES = ROOT / "results" / "figures"

OUT_DOCX = ROOT / "manuscript" / "manuscript_gigascience_ml_it.docx"
SUBMISSION_DIR = ROOT / "submission_gigascience"
SUBMISSION_FIG_DIR = SUBMISSION_DIR / "figures"


FIGURE_MAP = [
    ("fig1.png", "fig1_pipeline_flow.png", "End-to-end computational workflow."),
    ("fig2.png", "fig2_variant_counts_filtered.png", "Per-sample filtered variant counts."),
    ("fig3.png", "fig3_variant_counts_cohort_filtered.png", "Per-sample cohort-filtered variant counts."),
    ("fig4.png", "fig6_ml_repeatedcv_baseline.png", "Repeated cross-validation model comparison (filtered features)."),
    ("fig5.png", "fig7_ml_repeatedcv_cohort.png", "Repeated cross-validation model comparison (cohort-filtered features)."),
    ("fig6.png", "fig7_pca_filtered_class_donor.png", "PCA by class and donor for filtered features."),
    ("fig7.png", "fig8_pca_cohort_class_donor.png", "PCA by class and donor for cohort-filtered features."),
    ("fig8.png", "fig10_enrichment_top_terms.png", "Top enrichment terms summary."),
]


def set_document_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)


def add_page_numbers(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.text = "Page "
        run = p.add_run()

        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")

        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)


def add_par(
    doc: Document,
    text: str,
    align: WD_ALIGN_PARAGRAPH | None = None,
    bold: bool = False,
) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold


def add_table_from_df(doc: Document, title: str, df: pd.DataFrame) -> None:
    add_par(doc, title, bold=True)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            v = row[col]
            if isinstance(v, float):
                if abs(v) < 0.001 and v != 0:
                    text = f"{v:.2e}"
                else:
                    text = f"{v:.6f}".rstrip("0").rstrip(".")
                cells[i].text = text
            else:
                cells[i].text = str(v)
    doc.add_paragraph("")


def load_core_tables() -> dict[str, pd.DataFrame]:
    tables = {
        "table0": pd.read_csv(TABLES / "table0_external_metadata_geosra.csv"),
        "table1": pd.read_csv(TABLES / "table1_sample_summary.csv"),
        "table2": pd.read_csv(TABLES / "table2_variant_counts_summary.csv"),
        "table3": pd.read_csv(TABLES / "table3_ml_summary.csv"),
        "leakage": pd.read_csv(TABLES / "table_leakage_sensitivity.csv"),
        "metrics_repeated": pd.read_csv(TABLES / "supervised_metrics_repeated.csv"),
    }
    cmp_path = TABLES / "table_bcftools_vs_gatk_summary.csv"
    if cmp_path.exists():
        tables["caller_comparison"] = pd.read_csv(cmp_path)
    return tables


def extract_metrics(core: dict[str, pd.DataFrame]) -> dict[str, float | int]:
    t1 = core["table1"]
    t2 = core["table2"]
    t3 = core["table3"]
    leak = core["leakage"]
    t0 = core["table0"]

    n_total = int(t1["n_samples"].sum())
    n_uwe = int(t1.loc[t1["condition"] == "unwounded_skin", "n_samples"].iloc[0])
    n_we = int(t1.loc[t1["condition"] == "wound_edge", "n_samples"].iloc[0])
    n_donors = int(t0["donor_proxy_id"].nunique())

    t2f = t2.loc[t2["set"] == "filtered"].iloc[0]
    t2c = t2.loc[t2["set"] == "cohort_filtered"].iloc[0]
    t3f = t3.loc[t3["feature_set"] == "filtered"].iloc[0]
    t3c = t3.loc[t3["feature_set"] == "cohort_filtered"].iloc[0]
    leak_run = leak.loc[leak["evaluation"] == "run_level_repeated_stratified_5x30"].iloc[0]
    leak_donor = leak.loc[leak["evaluation"] == "donor_group_kfold_7"].iloc[0]

    return {
        "n_total": n_total,
        "n_uwe": n_uwe,
        "n_we": n_we,
        "n_donors": n_donors,
        "records_filtered_median": int(t2f["records_median"]),
        "records_filtered_min": int(t2f["records_min"]),
        "records_filtered_max": int(t2f["records_max"]),
        "records_cohort_median": int(t2c["records_median"]),
        "records_cohort_min": int(t2c["records_min"]),
        "records_cohort_max": int(t2c["records_max"]),
        "features_filtered": int(t3f["features"]),
        "features_cohort": int(t3c["features"]),
        "bacc_filtered": float(t3f["best_balanced_accuracy_mean"]),
        "bacc_cohort": float(t3c["best_balanced_accuracy_mean"]),
        "p_filtered": float(t3f["permutation_pvalue"]),
        "p_cohort": float(t3c["permutation_pvalue"]),
        "repeats": int(t3f["repeats"]),
        "permutations": int(t3f["permutations"]),
        "run_level_bacc": float(leak_run["balanced_accuracy_mean"]),
        "donor_level_bacc": float(leak_donor["balanced_accuracy_mean"]),
    }


def build_enrichment_table() -> pd.DataFrame:
    rows = []
    files = [
        ("top_mutated_genes", "enrichment_results_top_mutated_genes.csv"),
        ("wound_edge_higher_mutation_burden", "enrichment_results_wound_edge_higher_mutation_burden.csv"),
        ("unwounded_skin_higher_mutation_burden", "enrichment_results_unwounded_skin_higher_mutation_burden.csv"),
    ]
    for label, fn in files:
        df = pd.read_csv(TABLES / fn).sort_values("p_value").head(4)
        keep = df[["name", "source", "native", "p_value"]].copy()
        keep.insert(0, "gene_set", label)
        rows.append(keep)
    out = pd.concat(rows, ignore_index=True)
    return out


def to_display_samples(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out["condition"] = out["condition"].map(
        {"unwounded_skin": "Unwounded skin", "wound_edge": "Wound edge"}
    )
    return out.rename(columns={"condition": "Condition", "n_samples": "n"})


def to_display_variants(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out["set"] = out["set"].map({"filtered": "Filtered", "cohort_filtered": "Cohort-filtered"})
    return out.rename(
        columns={
            "set": "Feature set",
            "n_samples": "n",
            "records_median": "Records (median)",
            "records_min": "Records (min)",
            "records_max": "Records (max)",
            "snps_median": "SNPs (median)",
            "indels_median": "Indels (median)",
        }
    )


def to_display_ml(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out["feature_set"] = out["feature_set"].map(
        {"filtered": "Filtered", "cohort_filtered": "Cohort-filtered"}
    )
    return out.rename(
        columns={
            "feature_set": "Feature set",
            "samples": "Samples",
            "features": "Features",
            "best_balanced_accuracy_mean": "Best balanced accuracy",
            "permutation_pvalue": "Permutation p-value",
            "best_model_repeated": "Best repeated-CV model",
            "repeats": "Repeats",
            "permutations": "Permutations",
        }
    )


def to_display_leakage(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out["evaluation"] = out["evaluation"].map(
        {
            "run_level_repeated_stratified_5x30": "Run-level repeated stratified CV (5x30)",
            "donor_group_kfold_7": "Donor-grouped KFold (7 folds)",
        }
    )
    return out.rename(
        columns={
            "evaluation": "Evaluation protocol",
            "n_samples": "n_samples",
            "n_unique_donors": "n_donors",
            "accuracy_mean": "Accuracy mean",
            "balanced_accuracy_mean": "Balanced accuracy mean",
        }
    )


def to_top_models(df: pd.DataFrame) -> pd.DataFrame:
    out = (
        df.sort_values("balanced_accuracy_mean", ascending=False)
        .head(6)
        .loc[:, ["model", "balanced_accuracy_mean", "balanced_accuracy_std", "roc_auc_mean"]]
        .copy()
    )
    return out.rename(
        columns={
            "model": "Model",
            "balanced_accuracy_mean": "Balanced accuracy mean",
            "balanced_accuracy_std": "Balanced accuracy std",
            "roc_auc_mean": "ROC AUC mean",
        }
    )


def prepare_submission_figure_files() -> pd.DataFrame:
    SUBMISSION_FIG_DIR.mkdir(parents=True, exist_ok=True)
    rows: list[dict[str, str]] = []
    for target_name, source_name, caption in FIGURE_MAP:
        src = FIGURES / source_name
        dst = SUBMISSION_FIG_DIR / target_name
        if src.exists():
            copy2(src, dst)
            status = "copied"
        else:
            status = "missing_source"
        rows.append(
            {
                "Submission filename": target_name,
                "Source file": source_name,
                "Status": status,
                "Caption": caption,
            }
        )
    return pd.DataFrame(rows)


def main() -> int:
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    SUBMISSION_DIR.mkdir(parents=True, exist_ok=True)

    core = load_core_tables()
    m = extract_metrics(core)
    enrichment = build_enrichment_table()
    top_models = to_top_models(core["metrics_repeated"])
    fig_files = prepare_submission_figure_files()

    doc = Document()
    set_document_style(doc)
    add_page_numbers(doc)

    add_par(
        doc,
        "A reproducible ML pipeline for scRNA-seq variant-derived classification: wound edge versus unwounded skin",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
    )
    add_par(doc, "Kunikeyev Aidyn1,*", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_par(doc, "1 Affiliation: to be filled before submission", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_par(doc, "*Corresponding author email: to be filled; ORCID: to be filled", align=WD_ALIGN_PARAGRAPH.CENTER)

    add_par(doc, "Abstract", bold=True)
    add_par(
        doc,
        "Background: This work presents an IT/ML-focused, reproducible workflow that transforms scRNA-seq "
        "expressed-variant calls into sample-level classification features. "
        "Methods: Using PRJNA736095, we processed "
        f"{m['n_total']} runs ({m['n_uwe']} unwounded skin, {m['n_we']} wound edge; {m['n_donors']} donors) "
        "through alignment, variant calling, cohort-frequency filtering, variant-to-gene aggregation, repeated "
        "cross-validation, and permutation testing. "
        "Results: Median per-sample variant records were "
        f"{m['records_filtered_median']:,} for filtered VCFs and {m['records_cohort_median']:,} for cohort-filtered VCFs. "
        "Best repeated-CV balanced accuracy was "
        f"{m['bacc_filtered']:.3f} (filtered) and {m['bacc_cohort']:.3f} (cohort-filtered), with permutation p-values "
        f"{m['p_filtered']:.5f} and {m['p_cohort']:.5f}, respectively. "
        "A grouped-split sensitivity check showed balanced accuracy dropping from "
        f"{m['run_level_bacc']:.3f} (run-level CV) to {m['donor_level_bacc']:.3f} (donor-grouped CV). "
        "Conclusions: The main contribution is a practical reproducibility and validation framework for "
        "high-dimensional sparse omics features, with minimal and clearly separated biological interpretation."
    )

    add_par(
        doc,
        "Keywords: scRNA-seq, variant calling, feature engineering, machine learning, cross-validation, reproducibility",
    )

    add_par(doc, "1 Introduction", bold=True)
    add_par(
        doc,
        "The objective is to provide a reproducible machine-learning workflow for noisy, sparse, and "
        "high-dimensional sequencing-derived features. Rather than prioritizing biological novelty claims, "
        "the manuscript prioritizes auditable data transformations, robust model validation, and explicit "
        "checks for split-induced optimism."
    )

    add_par(doc, "2 Materials and Methods", bold=True)
    add_par(doc, "2.1 Dataset", bold=True)
    add_par(
        doc,
        "Public sequencing data were obtained from NCBI SRA BioProject PRJNA736095. "
        "The analyzed subset includes 14 runs with binary labels unwounded_skin and wound_edge."
    )
    add_par(doc, "2.2 Computational pipeline", bold=True)
    add_par(
        doc,
        "Reads were aligned with STARsolo; variants were called with GATK HaplotypeCaller; high-confidence records were "
        "retained using thresholds DP>=10, QUAL>=30, AD_alt>=3, VAF>=0.10; recurrent cohort positions "
        "(>=4 samples) were excluded for a cohort-filtered feature branch."
    )
    add_par(doc, "2.3 Feature engineering and ML", bold=True)
    add_par(
        doc,
        "Variants were mapped to exon-overlapping genes with BEDTools. Per-sample gene-level burden matrices "
        "were used for supervised learning (logistic regression, linear SVC, random forest, PCA-regularized and "
        "feature-selected variants) with repeated stratified cross-validation and permutation testing."
    )
    add_par(doc, "2.4 Minimal bioinformatics context", bold=True)
    add_par(
        doc,
        "A compact enrichment layer was retained only for contextual interpretation and not treated as the main "
        "scientific claim in this IT-focused manuscript."
    )

    add_par(doc, "3 Results", bold=True)
    add_par(doc, "3.1 Cohort and variant summary", bold=True)
    add_par(
        doc,
        f"Filtered records median/min/max: {m['records_filtered_median']:,}/"
        f"{m['records_filtered_min']:,}/{m['records_filtered_max']:,}. "
        f"Cohort-filtered records median/min/max: {m['records_cohort_median']:,}/"
        f"{m['records_cohort_min']:,}/{m['records_cohort_max']:,}."
    )
    add_table_from_df(doc, "Table 1. Class distribution.", to_display_samples(core["table1"]))
    add_table_from_df(doc, "Table 2. Variant count summary.", to_display_variants(core["table2"]))
    if "caller_comparison" in core:
        add_table_from_df(doc, "Table S1. Historical baseline vs active GATK comparison.", core["caller_comparison"])

    add_par(doc, "3.2 Classification performance", bold=True)
    add_par(
        doc,
        f"Feature dimensions were {m['features_filtered']:,} (filtered) and {m['features_cohort']:,} (cohort-filtered). "
        f"Best balanced accuracy was {m['bacc_filtered']:.3f} and {m['bacc_cohort']:.3f}, with permutation p-values "
        f"{m['p_filtered']:.5f} and {m['p_cohort']:.5f}."
    )
    add_table_from_df(doc, "Table 3. Main ML summary.", to_display_ml(core["table3"]))
    add_table_from_df(doc, "Table 4. Top repeated-CV models.", top_models)

    add_par(doc, "3.3 Split-sensitivity (leakage risk) analysis", bold=True)
    add_par(
        doc,
        f"Balanced accuracy changed from {m['run_level_bacc']:.3f} under run-level CV to "
        f"{m['donor_level_bacc']:.3f} under donor-grouped CV, indicating substantial dependence on "
        "split strategy in this small cohort."
    )
    add_table_from_df(doc, "Table 5. Leakage-sensitivity comparison.", to_display_leakage(core["leakage"]))

    add_par(doc, "3.4 Minimal enrichment summary", bold=True)
    add_table_from_df(
        doc,
        "Table 6. Top enrichment terms by gene-set category.",
        enrichment.rename(
            columns={
                "gene_set": "Gene set",
                "name": "Term",
                "source": "Source",
                "native": "Native ID",
                "p_value": "p-value",
            }
        ),
    )

    add_par(doc, "4 Discussion", bold=True)
    add_par(
        doc,
        "The key value is methodological: end-to-end reproducibility and explicit validation safeguards for "
        "small-sample omics classification. High run-level scores should be interpreted with caution unless "
        "grouped-split sensitivity is also reported."
    )

    add_par(doc, "5 Conclusions", bold=True)
    add_par(
        doc,
        "This manuscript provides an IT-ready and auditable framework for variant-derived ML classification "
        "from scRNA-seq and documents both performance and validation sensitivity."
    )

    add_par(doc, "List of abbreviations", bold=True)
    abbreviations = [
        "AD: alternate allele depth",
        "CV: cross-validation",
        "DP: read depth",
        "ML: machine learning",
        "PCA: principal component analysis",
        "SRA: Sequence Read Archive",
        "VAF: variant allele fraction",
    ]
    for item in sorted(abbreviations):
        add_par(doc, item)

    add_par(doc, "Figure legends (figures uploaded as separate files)", bold=True)
    for i, (submission_name, _, caption) in enumerate(FIGURE_MAP, start=1):
        add_par(doc, f"Figure {i} ({submission_name}). {caption}")

    add_table_from_df(doc, "Figure file checklist for submission.", fig_files)

    add_par(doc, "Declarations", bold=True)
    add_par(doc, "Ethics approval and consent to participate: Not applicable.")
    add_par(doc, "Consent for publication: Not applicable.")
    add_par(
        doc,
        "Data Availability: The sequencing dataset supporting this article is publicly available at NCBI SRA "
        "under BioProject accession PRJNA736095. Run-level accessions are listed in "
        "Final version/results/tables/table0_external_metadata_geosra.csv. "
        "Processed outputs used in this manuscript are available in Final version/results/. "
        "A public archive DOI for processed artifacts/code should be added before submission if available.",
    )
    add_par(doc, "Competing interests: KA declares that there are no competing interests.")
    add_par(
        doc,
        "Funding: Not applicable (no funding information was found in workspace files).",
    )
    add_par(
        doc,
        "Authors' contributions: KA designed the pipeline, implemented analyses, interpreted outputs, and wrote the manuscript.",
    )
    add_par(
        doc,
        "Acknowledgements: Public data providers and original PRJNA736095 contributors are acknowledged.",
    )
    add_par(
        doc,
        "Authors' information (optional): KA is an independent researcher; full institutional details to be completed before submission.",
    )

    add_par(doc, "Availability and requirements", bold=True)
    add_par(doc, "Project name: somatic-like-scrna-variants-ml-pipeline")
    add_par(doc, "Project home page: https://github.com/dauletulyaidyn/somatic-like-scrna-variants-ml-pipeline")
    add_par(doc, "Operating system(s): Linux/WSL; Windows host paths used in this workspace")
    add_par(doc, "Programming language: Python")
    add_par(doc, "Other requirements: STAR/STARsolo, samtools, GATK 4, bedtools, scikit-learn")
    add_par(doc, "License: to be confirmed from repository metadata before submission")
    add_par(doc, "RRID: Not applicable")
    add_par(doc, "bio.tools ID: Not applicable")

    add_par(doc, "Disclosure of use of AI-assisted tools including generative AI", bold=True)
    add_par(
        doc,
        "A generative AI assistant was used for editorial drafting support during manuscript preparation. "
        "All numerical results, interpretations, and final wording were reviewed and verified by the author. "
        "No AI-assisted tool was used to create or modify scientific figures.",
    )

    add_par(doc, "References", bold=True)
    refs = [
        "[1] Dobin A, Davis CA, Schlesinger F, et al. STAR: ultrafast universal RNA-seq aligner. Bioinformatics. 2013;29(1):15-21.",
        "[2] Li H, Handsaker B, Wysoker A, et al. The Sequence Alignment/Map format and SAMtools. Bioinformatics. 2009;25(16):2078-2079.",
        "[3] Danecek P, Bonfield JK, Liddle J, et al. Twelve years of SAMtools and BCFtools. GigaScience. 2021;10(2):giab008.",
        "[4] Quinlan AR, Hall IM. BEDTools: a flexible suite of utilities for comparing genomic features. Bioinformatics. 2010;26(6):841-842.",
        "[5] Pedregosa F, Varoquaux G, Gramfort A, et al. Scikit-learn: Machine Learning in Python. J Mach Learn Res. 2011;12:2825-2830.",
        "[6] Ojala M, Garriga GC. Permutation Tests for Studying Classifier Performance. J Mach Learn Res. 2010;11:1833-1863.",
        "[7] Wolf FA, Angerer P, Theis FJ. SCANPY: large-scale single-cell gene expression data analysis. Genome Biol. 2018;19:15.",
        "[8] Traag VA, Waltman L, van Eck NJ. From Louvain to Leiden: guaranteeing well-connected communities. Sci Rep. 2019;9:5233.",
        "[9] McInnes L, Healy J, Melville J. UMAP: Uniform Manifold Approximation and Projection for Dimension Reduction. arXiv:1802.03426. Accessed 2026-02-22.",
        "[10] Huang X, Huang Y. Cellsnp-lite: an efficient tool for genotyping single cells. Bioinformatics. 2021;37(23):4569-4571.",
        "[11] Raudvere U, Kolberg L, Kuzmin I, et al. g:Profiler: a web server for functional enrichment analysis and conversions of gene lists (2019 update). Nucleic Acids Res. 2019;47(W1):W191-W198.",
        "[12] NCBI BioProject. PRJNA736095. https://www.ncbi.nlm.nih.gov/bioproject/PRJNA736095. Accessed 2026-02-22.",
    ]
    for ref in refs:
        add_par(doc, ref)

    doc.save(str(OUT_DOCX))
    print(f"Wrote manuscript: {OUT_DOCX}")
    print(f"Prepared figure files: {SUBMISSION_FIG_DIR}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
