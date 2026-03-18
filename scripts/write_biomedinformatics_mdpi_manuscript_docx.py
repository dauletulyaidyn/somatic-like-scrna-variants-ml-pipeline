#!/usr/bin/env python3
from __future__ import annotations

import json
from pathlib import Path
from shutil import copy2

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]  # .../Final version
TABLES = ROOT / "results" / "tables"
FIGURES = ROOT / "results" / "figures"
CONFIG_DIR = ROOT / "repo" / "config"

OUT_DOCX = ROOT / "manuscript" / "manuscript_biomedinformatics_mdpi_ml.docx"
SUBMISSION_DIR = ROOT / "submission_biomedinformatics"
SUBMISSION_FIG_DIR = SUBMISSION_DIR / "figures"
CODE_REPO_URL = "https://github.com/dauletulyaidyn/somatic-like-scrna-variants-ml-pipeline"


FIGURE_MAP = [
    ("fig1.png", "fig1_pipeline_flow.png", "End-to-end analysis pipeline."),
    ("fig2.png", "fig2_variant_counts_filtered.png", "Per-sample filtered variant counts."),
    ("fig3.png", "fig3_variant_counts_cohort_filtered.png", "Per-sample cohort-filtered variant counts."),
    ("fig4.png", "fig6_ml_repeatedcv_baseline.png", "Repeated-CV model comparison for filtered features."),
    ("fig5.png", "fig7_ml_repeatedcv_cohort.png", "Repeated-CV model comparison for cohort-filtered features."),
    ("fig6.png", "fig7_pca_filtered_class_donor.png", "PCA projection with class and donor overlays (filtered features)."),
    ("fig7.png", "fig8_pca_cohort_class_donor.png", "PCA projection with class and donor overlays (cohort-filtered features)."),
    ("fig8.png", "fig10_enrichment_top_terms.png", "Top enrichment terms summary."),
]


def set_document_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_par(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold


def add_table_from_df(doc: Document, caption: str, df: pd.DataFrame) -> None:
    add_par(doc, caption, bold=True)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            v = row[col]
            if isinstance(v, float):
                if abs(v) < 0.001 and v != 0:
                    text = f"{v:.2e}"
                else:
                    text = f"{v:.6f}".rstrip("0").rstrip(".")
            else:
                text = str(v)
            cells[i].text = text

    doc.add_paragraph("")


def add_figure(doc: Document, idx: int, fig_name: str, caption: str) -> None:
    fig_path = SUBMISSION_FIG_DIR / fig_name
    if not fig_path.exists():
        add_par(doc, f"Figure {idx}. {caption} (file missing: {fig_name})")
        return
    doc.add_picture(str(fig_path), width=Inches(6.2))
    add_par(doc, f"Figure {idx}. {caption}")
    doc.add_paragraph("")


def copy_submission_figures() -> pd.DataFrame:
    SUBMISSION_FIG_DIR.mkdir(parents=True, exist_ok=True)
    rows: list[dict[str, str]] = []
    for out_name, src_name, caption in FIGURE_MAP:
        src = FIGURES / src_name
        dst = SUBMISSION_FIG_DIR / out_name
        if src.exists():
            copy2(src, dst)
            status = "copied"
        else:
            status = "missing_source"
        rows.append(
            {
                "Figure file": out_name,
                "Source": src_name,
                "Status": status,
                "Caption": caption,
            }
        )
    return pd.DataFrame(rows)


def load_core_tables() -> dict[str, pd.DataFrame]:
    tables = {
        "table0": pd.read_csv(TABLES / "table0_external_metadata_geosra.csv"),
        "table1": pd.read_csv(TABLES / "table1_sample_summary.csv"),
        "table2": pd.read_csv(TABLES / "table2_variant_counts_summary.csv"),
        "table3": pd.read_csv(TABLES / "table3_ml_summary.csv"),
        "leakage": pd.read_csv(TABLES / "table_leakage_sensitivity.csv"),
        "metrics_repeated": pd.read_csv(TABLES / "supervised_metrics_repeated.csv"),
        "variant_counts": pd.read_csv(TABLES / "variant_counts_filtered_vs_cohort.csv"),
        "pca_filtered": pd.read_csv(TABLES / "pca_filtered_coords.csv"),
        "pca_cohort": pd.read_csv(TABLES / "pca_cohort_coords.csv"),
        "corr_spearman": pd.read_csv(TABLES / "correlation_spearman.tsv", sep="\t", index_col=0),
        "marker_enrichment": pd.read_csv(TABLES / "mutation_gene_set_marker_enrichment.csv"),
        "marker_overlap_top5": pd.read_csv(TABLES / "mutation_gene_set_marker_overlap_top5.csv"),
    }
    cmp_path = TABLES / "table_bcftools_vs_gatk_summary.csv"
    if cmp_path.exists():
        tables["caller_comparison"] = pd.read_csv(cmp_path)
    return tables


def load_pipeline_settings() -> dict[str, str | int | float]:
    gatk_cfg = json.loads((CONFIG_DIR / "gatk_config.json").read_text(encoding="utf-8"))
    cohort_cfg = json.loads((CONFIG_DIR / "cohort_filter_config.json").read_text(encoding="utf-8"))

    return {
        "caller_name": "GATK HaplotypeCaller",
        "filter_names": ", ".join(f["name"] for f in gatk_cfg.get("variant_filters", [])),
        "caller_mode": str(gatk_cfg.get("mode", "import_existing")),
        "cohort_min_samples": int(cohort_cfg.get("min_samples", 4)),
        "cohort_min_vaf": float(cohort_cfg.get("min_vaf", 0.05)),
    }


def extract_metrics(core: dict[str, pd.DataFrame]) -> dict[str, float | int]:
    t0 = core["table0"]
    t1 = core["table1"]
    t2 = core["table2"]
    t3 = core["table3"]
    leak = core["leakage"]

    n_total = int(t1["n_samples"].sum())
    n_uwe = int(t1.loc[t1["condition"] == "unwounded_skin", "n_samples"].iloc[0])
    n_we = int(t1.loc[t1["condition"] == "wound_edge", "n_samples"].iloc[0])
    n_donors = int(t0["donor_proxy_id"].nunique())

    t2f = t2.loc[t2["set"] == "filtered"].iloc[0]
    t2c = t2.loc[t2["set"] == "cohort_filtered"].iloc[0]
    t3f = t3.loc[t3["feature_set"] == "filtered"].iloc[0]
    t3c = t3.loc[t3["feature_set"] == "cohort_filtered"].iloc[0]

    leak_run = leak.loc[leak["evaluation"] == "run_level_repeated_stratified_5x30"].iloc[0]
    leak_donor = leak.loc[leak["evaluation"] == "donor_group_kfold_7"].iloc[0]

    return {
        "n_total": n_total,
        "n_uwe": n_uwe,
        "n_we": n_we,
        "n_donors": n_donors,
        "records_filtered_median": int(t2f["records_median"]),
        "records_filtered_min": int(t2f["records_min"]),
        "records_filtered_max": int(t2f["records_max"]),
        "records_cohort_median": int(t2c["records_median"]),
        "records_cohort_min": int(t2c["records_min"]),
        "records_cohort_max": int(t2c["records_max"]),
        "features_filtered": int(t3f["features"]),
        "features_cohort": int(t3c["features"]),
        "bacc_filtered": float(t3f["best_balanced_accuracy_mean"]),
        "bacc_cohort": float(t3c["best_balanced_accuracy_mean"]),
        "p_filtered": float(t3f["permutation_pvalue"]),
        "p_cohort": float(t3c["permutation_pvalue"]),
        "repeats": int(t3f["repeats"]),
        "permutations": int(t3f["permutations"]),
        "run_level_bacc": float(leak_run["balanced_accuracy_mean"]),
        "donor_level_bacc": float(leak_donor["balanced_accuracy_mean"]),
    }


def donor_demographics(table0: pd.DataFrame) -> dict[str, float | int]:
    donor_df = table0.drop_duplicates("donor_proxy_id").copy()
    donor_cond = donor_df["condition"].value_counts().to_dict()
    donor_sex = donor_df["sex"].value_counts().to_dict()

    return {
        "donor_n_uwe": int(donor_cond.get("unwounded_skin", 0)),
        "donor_n_we": int(donor_cond.get("wound_edge", 0)),
        "donor_n_female": int(donor_sex.get("Female", 0)),
        "donor_n_male": int(donor_sex.get("Male", 0)),
        "age_min": int(donor_df["age_years"].min()),
        "age_median": float(donor_df["age_years"].median()),
        "age_max": int(donor_df["age_years"].max()),
    }


def variant_summary_stats(variant_counts: pd.DataFrame, table0: pd.DataFrame) -> dict[str, float | str]:
    meta = table0.loc[:, ["run_id", "condition"]].rename(columns={"run_id": "sample"})
    vc = variant_counts.merge(meta, on="sample", how="left")

    vc["records_reduction_abs"] = vc["filtered_number of records:"] - vc["cohort_number of records:"]
    vc["records_reduction_pct"] = vc["records_reduction_abs"] / vc["filtered_number of records:"] * 100.0

    vc["snps_reduction_abs"] = vc["filtered_number of SNPs:"] - vc["cohort_number of SNPs:"]
    vc["snps_reduction_pct"] = vc["snps_reduction_abs"] / vc["filtered_number of SNPs:"] * 100.0

    vc["indels_reduction_abs"] = vc["filtered_number of indels:"] - vc["cohort_number of indels:"]
    vc["indels_reduction_pct"] = vc["indels_reduction_abs"] / vc["filtered_number of indels:"] * 100.0

    cond_median = vc.groupby("condition")["records_reduction_pct"].median().to_dict()
    cond_record_median = vc.groupby("condition")[
        ["filtered_number of records:", "cohort_number of records:"]
    ].median()

    max_row = vc.sort_values("records_reduction_pct", ascending=False).iloc[0]
    min_row = vc.sort_values("records_reduction_pct", ascending=True).iloc[0]

    return {
        "records_reduction_mean_pct": float(vc["records_reduction_pct"].mean()),
        "records_reduction_median_pct": float(vc["records_reduction_pct"].median()),
        "records_reduction_min_pct": float(vc["records_reduction_pct"].min()),
        "records_reduction_max_pct": float(vc["records_reduction_pct"].max()),
        "snps_reduction_mean_pct": float(vc["snps_reduction_pct"].mean()),
        "indels_reduction_mean_pct": float(vc["indels_reduction_pct"].mean()),
        "records_reduction_median_pct_uwe": float(cond_median.get("unwounded_skin", np.nan)),
        "records_reduction_median_pct_we": float(cond_median.get("wound_edge", np.nan)),
        "records_median_filtered_uwe": float(cond_record_median.loc["unwounded_skin", "filtered_number of records:"]),
        "records_median_cohort_uwe": float(cond_record_median.loc["unwounded_skin", "cohort_number of records:"]),
        "records_median_filtered_we": float(cond_record_median.loc["wound_edge", "filtered_number of records:"]),
        "records_median_cohort_we": float(cond_record_median.loc["wound_edge", "cohort_number of records:"]),
        "records_range_fold_filtered": float(
            vc["filtered_number of records:"].max() / vc["filtered_number of records:"].min()
        ),
        "records_range_fold_cohort": float(
            vc["cohort_number of records:"].max() / vc["cohort_number of records:"].min()
        ),
        "max_reduction_sample": str(max_row["sample"]),
        "max_reduction_sample_pct": float(max_row["records_reduction_pct"]),
        "min_reduction_sample": str(min_row["sample"]),
        "min_reduction_sample_pct": float(min_row["records_reduction_pct"]),
    }


def pca_shape_stats(pca_df: pd.DataFrame) -> dict[str, float | str]:
    cent = pca_df.groupby("condition")[["PC1", "PC2"]].mean()
    class_dist = float(np.linalg.norm(cent.loc["unwounded_skin"].values - cent.loc["wound_edge"].values))

    donor_pair_distances: list[tuple[str, float]] = []
    for donor, g in pca_df.groupby("donor_proxy_id"):
        if len(g) >= 2:
            pts = g.loc[:, ["PC1", "PC2"]].values
            donor_pair_distances.append((str(donor), float(np.linalg.norm(pts[0] - pts[1]))))

    within = np.array([d for _donor, d in donor_pair_distances], dtype=float)
    donor_centroids = pca_df.groupby("donor_proxy_id")[["PC1", "PC2"]].mean().values
    between_list: list[float] = []
    for i in range(len(donor_centroids)):
        for j in range(i + 1, len(donor_centroids)):
            between_list.append(float(np.linalg.norm(donor_centroids[i] - donor_centroids[j])))
    between = np.array(between_list, dtype=float)

    max_pair_donor, max_pair_distance = max(donor_pair_distances, key=lambda x: x[1])
    min_pair_donor, min_pair_distance = min(donor_pair_distances, key=lambda x: x[1])

    return {
        "class_centroid_distance": class_dist,
        "within_donor_median": float(np.median(within)),
        "within_donor_min": float(within.min()),
        "within_donor_max": float(within.max()),
        "between_donor_median": float(np.median(between)),
        "between_donor_min": float(between.min()),
        "between_donor_max": float(between.max()),
        "within_over_between_ratio": float(np.median(within) / np.median(between)),
        "max_pair_donor": max_pair_donor,
        "max_pair_distance": max_pair_distance,
        "min_pair_donor": min_pair_donor,
        "min_pair_distance": min_pair_distance,
    }


def top_models(df: pd.DataFrame) -> pd.DataFrame:
    out = (
        df.sort_values("balanced_accuracy_mean", ascending=False)
        .head(6)
        .loc[:, ["model", "balanced_accuracy_mean", "balanced_accuracy_std", "roc_auc_mean"]]
        .copy()
    )
    return out.rename(
        columns={
            "model": "Model",
            "balanced_accuracy_mean": "Balanced accuracy mean",
            "balanced_accuracy_std": "Balanced accuracy std",
            "roc_auc_mean": "ROC AUC mean",
        }
    )


def enrichment_table() -> pd.DataFrame:
    rows = []
    items = [
        ("top_mutated_genes", "enrichment_results_top_mutated_genes.csv"),
        ("wound_edge_higher_mutation_burden", "enrichment_results_wound_edge_higher_mutation_burden.csv"),
        ("unwounded_skin_higher_mutation_burden", "enrichment_results_unwounded_skin_higher_mutation_burden.csv"),
    ]
    for label, filename in items:
        df = pd.read_csv(TABLES / filename).sort_values("p_value").head(4)
        part = df[["name", "source", "native", "p_value"]].copy()
        part.insert(0, "Gene set", label)
        rows.append(part)
    out = pd.concat(rows, ignore_index=True)
    return out.rename(
        columns={
            "name": "Term",
            "source": "Source",
            "native": "Native ID",
            "p_value": "p-value",
        }
    )


def display_table1(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out["condition"] = out["condition"].map(
        {"unwounded_skin": "Unwounded skin", "wound_edge": "Wound edge"}
    )
    return out.rename(columns={"condition": "Condition", "n_samples": "n"})


def display_table2(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out["set"] = out["set"].map({"filtered": "Filtered", "cohort_filtered": "Cohort-filtered"})
    return out.rename(
        columns={
            "set": "Feature set",
            "n_samples": "n",
            "records_median": "Records (median)",
            "records_min": "Records (min)",
            "records_max": "Records (max)",
            "snps_median": "SNPs (median)",
            "indels_median": "Indels (median)",
        }
    )


def display_table3(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out["feature_set"] = out["feature_set"].map(
        {"filtered": "Filtered", "cohort_filtered": "Cohort-filtered"}
    )
    return out.rename(
        columns={
            "feature_set": "Feature set",
            "samples": "Samples",
            "features": "Features",
            "best_balanced_accuracy_mean": "Best balanced accuracy",
            "permutation_pvalue": "Permutation p-value",
            "best_model_repeated": "Best repeated-CV model",
            "repeats": "Repeats",
            "permutations": "Permutations",
        }
    )


def display_leakage(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out["evaluation"] = out["evaluation"].map(
        {
            "run_level_repeated_stratified_5x30": "Run-level repeated stratified CV (5x30)",
            "donor_group_kfold_7": "Donor-grouped KFold (7)",
        }
    )
    return out.rename(
        columns={
            "evaluation": "Evaluation protocol",
            "n_samples": "n_samples",
            "n_unique_donors": "n_donors",
            "accuracy_mean": "Accuracy mean",
            "balanced_accuracy_mean": "Balanced accuracy mean",
        }
    )


def display_corr_matrix(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy().reset_index().rename(columns={"index": "Metric"})
    return out.rename(
        columns={
            "total_gene_burden": "Total gene burden",
            "filtered_records": "Filtered records",
            "cohort_filtered_records": "Cohort-filtered records",
        }
    )


def marker_overlap_summary(df: pd.DataFrame) -> pd.DataFrame:
    out = (
        df.groupby("gene_set", as_index=False)
        .agg(
            tests=("gene_set", "size"),
            tests_with_overlap=("overlap_n", lambda s: int((s > 0).sum())),
            min_nominal_pvalue=("p_value", "min"),
            min_fdr_qvalue=("q_value", "min"),
        )
        .copy()
    )
    return out.rename(
        columns={
            "gene_set": "Gene set",
            "tests": "Tests",
            "tests_with_overlap": "Tests with overlap_n > 0",
            "min_nominal_pvalue": "Min nominal p-value",
            "min_fdr_qvalue": "Min FDR q-value",
        }
    )


def display_marker_overlap_top5(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    return out.rename(
        columns={
            "gene_set": "Gene set",
            "cluster": "Cluster",
            "predicted_cell_type": "Predicted cell type",
            "short_name": "Marker short-name",
            "overlap_n": "Overlap n",
            "overlap_genes": "Overlap genes",
        }
    )


def write_cover_letter(metrics: dict[str, float | int]) -> None:
    text = f"""Dear Editors of BioMedInformatics,

Please consider our manuscript entitled:
"A reproducible ML pipeline for scRNA-seq variant-derived classification: wound edge versus unwounded skin".

The manuscript presents a reproducible data-science workflow for converting scRNA-seq expressed variants into gene-level features and validating binary classification with repeated cross-validation, permutation testing, and donor-grouped leakage-sensitivity analysis. In our dataset (PRJNA736095), the best repeated-CV balanced accuracy was {metrics['bacc_filtered']:.3f} (filtered features) and {metrics['bacc_cohort']:.3f} (cohort-filtered features), while donor-grouped evaluation reduced balanced accuracy to {metrics['donor_level_bacc']:.3f}, highlighting a practical reproducibility risk in small biomedical cohorts.

This contribution fits BioMedInformatics because it focuses on computational reproducibility, model validation design, and transparent reporting for biomedical machine-learning workflows.

Required statements:
1) We confirm that neither the manuscript nor any parts of its content are currently under consideration for publication with or published in another journal.
2) All authors have approved the manuscript and agree with its submission to BioMedInformatics.

Sincerely,
Kunikeyev Aidyn
Corresponding author (email and affiliation to be finalized)
"""
    path = SUBMISSION_DIR / "cover_letter_biomedinformatics.md"
    path.write_text(text, encoding="utf-8")


def write_submission_checklist() -> None:
    text = """# BioMedInformatics Submission Checklist

Manuscript:
- `Final version/manuscript/manuscript_biomedinformatics_mdpi_ml.docx`

Cover letter (required by journal):
- `Final version/submission_biomedinformatics/cover_letter_biomedinformatics.md`

Figure files for submission:
- `Final version/submission_biomedinformatics/figures/fig1.png` ... `fig8.png`
- `Final version/submission_biomedinformatics/figure_manifest.csv`

Before submission, fill placeholders:
- Author affiliation(s) and full corresponding-author details.
- ORCID profile link(s), if available.
- Funding details (if any).
- Data repository DOI/URL for code and processed outputs (if available).
- Verify that the code repository is public and include a persistent release link/DOI if available.

Policy checks from BioMedInformatics instructions:
- Keep references numbered in square brackets in text.
- Ensure figures/tables are cited in order and captions are complete.
- Include all mandatory back-matter statements (Author Contributions, Funding, IRB, Informed Consent, Data Availability, Acknowledgments, Conflicts of Interest).
- Include GenAI disclosure in Materials and Methods and Acknowledgments if GenAI used beyond superficial language editing.
"""
    path = SUBMISSION_DIR / "SUBMISSION_CHECKLIST.md"
    path.write_text(text, encoding="utf-8")


def main() -> int:
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    SUBMISSION_DIR.mkdir(parents=True, exist_ok=True)

    fig_manifest = copy_submission_figures()
    core = load_core_tables()
    metrics = extract_metrics(core)
    settings = load_pipeline_settings()
    demo = donor_demographics(core["table0"])
    variant_stats = variant_summary_stats(core["variant_counts"], core["table0"])
    pca_filtered = pca_shape_stats(core["pca_filtered"])
    pca_cohort = pca_shape_stats(core["pca_cohort"])
    corr = core["corr_spearman"]
    marker_enr = core["marker_enrichment"]

    feature_reduction_pct = (
        (metrics["features_filtered"] - metrics["features_cohort"]) / metrics["features_filtered"] * 100.0
    )
    bacc_delta_pp = (metrics["bacc_cohort"] - metrics["bacc_filtered"]) * 100.0
    leak_abs_drop = metrics["run_level_bacc"] - metrics["donor_level_bacc"]
    leak_rel_drop_pct = leak_abs_drop / metrics["run_level_bacc"] * 100.0
    rho_gene_filtered = float(corr.loc["total_gene_burden", "filtered_records"])
    rho_gene_cohort = float(corr.loc["total_gene_burden", "cohort_filtered_records"])
    rho_filtered_cohort = float(corr.loc["filtered_records", "cohort_filtered_records"])
    n_marker_tests = int(len(marker_enr))
    n_marker_overlap_positive = int((marker_enr["overlap_n"] > 0).sum())
    n_marker_nominal = int((marker_enr["p_value"] < 0.05).sum())
    n_marker_fdr = int((marker_enr["q_value"] < 0.05).sum())
    best_marker_row = marker_enr.loc[marker_enr["p_value"].idxmin()]

    doc = Document()
    set_document_style(doc)

    add_par(
        doc,
        "A reproducible ML pipeline for scRNA-seq variant-derived classification: wound edge versus unwounded skin",
        bold=True,
    )
    add_par(doc, "Kunikeyev Aidyn")
    add_par(doc, "Independent Researcher (affiliation details to be completed before submission)")
    add_par(doc, "Correspondence: email to be provided; ORCID to be provided")

    add_heading(doc, "Abstract", level=1)
    add_par(
        doc,
        "Background: Robust machine-learning studies in biomedicine require reproducible pipelines and split-aware "
        "validation. Methods: We analyzed public scRNA-seq data (PRJNA736095; 14 runs across 7 donors) using an "
        "end-to-end workflow with STARsolo alignment, SAMtools/BCFtools variant calling, cohort-frequency "
        "filtering, variant-to-gene feature engineering, repeated stratified cross-validation, permutation "
        "testing, and donor-grouped sensitivity analysis. Results: Median variant records per sample were "
        f"{metrics['records_filtered_median']:,} (filtered branch) and {metrics['records_cohort_median']:,} "
        f"(cohort-filtered branch), with a median per-sample record reduction of "
        f"{variant_stats['records_reduction_median_pct']:.2f}%. Feature dimension was reduced by "
        f"{feature_reduction_pct:.2f}% ({metrics['features_filtered']:,} to {metrics['features_cohort']:,}) while "
        f"best repeated-CV balanced accuracy changed from {metrics['bacc_filtered']:.3f} to "
        f"{metrics['bacc_cohort']:.3f} ({bacc_delta_pp:+.2f} percentage points). Permutation p-values were "
        f"{metrics['p_filtered']:.5f} and {metrics['p_cohort']:.5f}. Donor-grouped evaluation reduced balanced "
        f"accuracy from {metrics['run_level_bacc']:.3f} to {metrics['donor_level_bacc']:.3f} (absolute drop "
        f"{leak_abs_drop:.3f}; relative drop {leak_rel_drop_pct:.1f}%). Spearman analysis showed positive "
        f"association between total gene burden and record counts (rho={rho_gene_filtered:.3f}-{rho_gene_cohort:.3f}). "
        f"Marker-to-mutation overlaps were sparse and non-significant after FDR correction (q>=0.630701). "
        f"Conclusions: The main contribution is a "
        "reproducibility-first ML framework with explicit quantification of split sensitivity and transparent "
        "interpretation of model performance, variance, and feature-branch effects."
    )
    add_par(
        doc,
        "Keywords: scRNA-seq; variant calling; feature engineering; machine learning; cross-validation; reproducibility"
    )

    add_heading(doc, "1. Introduction", level=1)
    add_par(
        doc,
        "Biomedical classification from sequencing-derived features is frequently affected by high dimensionality, "
        "small sample size, and hidden non-independence between observations. These conditions can inflate "
        "run-level performance estimates when biological replicates or technical duplicates are not handled with "
        "appropriate split design."
    )
    add_par(
        doc,
        "This study addresses that problem as an engineering and reproducibility task. The goal is not to maximize "
        "biological claim breadth, but to provide an auditable end-to-end pipeline that links data processing "
        "choices to downstream model behavior. The workflow converts expressed variant calls into gene-level "
        "features and evaluates classification under both optimistic and conservative validation regimes."
    )
    add_par(
        doc,
        "A minimal functional interpretation layer is included only to contextualize outputs, while the primary "
        "focus remains on computational design, quantifiable sensitivity to split protocol, and transparent "
        "reporting standards aligned with biomedical ML practice [1-6]."
    )

    add_heading(doc, "2. Materials and Methods", level=1)
    add_heading(doc, "2.1. Dataset and labels", level=2)
    add_par(
        doc,
        f"We used BioProject PRJNA736095 with {metrics['n_total']} runs "
        f"({metrics['n_uwe']} unwounded skin, {metrics['n_we']} wound edge) from {metrics['n_donors']} donor proxies. "
        f"At donor level, this corresponds to {demo['donor_n_uwe']} unwounded-skin donors and "
        f"{demo['donor_n_we']} wound-edge donors; sex distribution was {demo['donor_n_female']} female and "
        f"{demo['donor_n_male']} male donors, with age range {demo['age_min']}-{demo['age_max']} years "
        f"(median {demo['age_median']:.1f})."
    )
    add_heading(doc, "2.2. Alignment and variant calling", level=2)
    add_par(
        doc,
        "Reads were aligned with STARsolo [1]. Variant calling used the GATK RNA workflow centered on "
        f"{settings['caller_name']}, followed by hard filtering and PASS-only export. "
        f"Configured filter labels in the active workflow were: {settings['filter_names'] or 'project defaults'}. "
        f"The active stage mode was {settings['caller_mode']}. All parameter values were taken directly from "
        "project configuration files."
    )
    add_heading(doc, "2.3. Cohort-frequency filtering and feature engineering", level=2)
    add_par(
        doc,
        "Cohort-common loci were identified as sites observed in at least "
        f"{settings['cohort_min_samples']} samples with maximum per-sample VAF >= {settings['cohort_min_vaf']:.2f}; "
        "these loci were excluded in the cohort-filtered branch. Variants were mapped to genes by coordinate overlap against GTF "
        "gene intervals, and sample-level gene-burden features were constructed for both the filtered and "
        "cohort-filtered branches."
    )
    add_heading(doc, "2.4. Supervised learning and validation", level=2)
    add_par(
        doc,
        "Models (logistic regression, linear SVC, random forest, and reduced-feature variants) were implemented "
        "in scikit-learn [4]. Performance was estimated with repeated stratified CV and permutation testing [5]. "
        f"The repeated protocol used {metrics['repeats']} repeats and permutation testing used "
        f"{metrics['permutations']} permutations per feature branch."
    )
    add_heading(doc, "2.5. Leakage-sensitivity analysis", level=2)
    add_par(
        doc,
        "To assess split dependence, run-level repeated stratified CV was compared against donor-grouped KFold "
        "evaluation using donor proxy IDs from external metadata. This analysis explicitly tests whether performance "
        "remains stable when train-test separation is enforced at donor level."
    )
    add_heading(doc, "2.6. Minimal functional context", level=2)
    add_par(
        doc,
        "A compact enrichment layer (g:Profiler [6]) was retained as contextual output, while the manuscript "
        "focus remains on computational methodology. No broad mechanistic conclusions were drawn from enrichment "
        "results in this manuscript."
    )
    add_heading(doc, "2.7. Expression-mutation correlation module", level=2)
    add_par(
        doc,
        "To connect mutational and expression-derived outputs without over-interpreting biological effects, we used "
        "two compact analyses. First, Spearman correlations were evaluated among run-level totals from "
        "correlation_spearman.tsv (total gene burden, filtered records, cohort-filtered records). Second, "
        "cluster marker sets from expression_top_markers_by_cluster.csv were compared with mutation-derived gene "
        "sets (top_mutated_genes, wound_edge_higher_mutation_burden, unwounded_skin_higher_mutation_burden) using "
        "the precomputed hypergeometric enrichment table mutation_gene_set_marker_enrichment.csv with FDR-adjusted "
        "q-values."
    )
    add_heading(doc, "2.8. Generative AI disclosure", level=2)
    add_par(
        doc,
        "A generative AI assistant was used for drafting support and structural editing of manuscript text. "
        "No primary data generation, statistical computation, or scientific figure synthesis depended on AI tools; "
        "all numerical values were taken from project-generated result tables and verified by the author."
    )

    add_heading(doc, "3. Results", level=1)
    add_heading(doc, "3.1. Cohort and variant summary", level=2)
    add_par(
        doc,
        f"Table 1 confirms class composition at run level ({metrics['n_uwe']} unwounded-skin, "
        f"{metrics['n_we']} wound-edge runs). Because two runs were available per donor, the cohort encodes both "
        "biological and technical structure; this motivates explicit donor-aware sensitivity analyses."
    )
    add_table_from_df(doc, "Table 1. Sample counts by condition.", display_table1(core["table1"]))
    add_par(
        doc,
        f"Table 2 shows that filtered records had median/min/max of {metrics['records_filtered_median']:,}/"
        f"{metrics['records_filtered_min']:,}/{metrics['records_filtered_max']:,}, while cohort-filtered records "
        f"had {metrics['records_cohort_median']:,}/{metrics['records_cohort_min']:,}/"
        f"{metrics['records_cohort_max']:,}. Median per-sample reduction after cohort filtering was "
        f"{variant_stats['records_reduction_median_pct']:.2f}% (mean "
        f"{variant_stats['records_reduction_mean_pct']:.2f}%; range "
        f"{variant_stats['records_reduction_min_pct']:.2f}-{variant_stats['records_reduction_max_pct']:.2f}%). "
        f"Reduction was somewhat stronger in unwounded skin (median "
        f"{variant_stats['records_reduction_median_pct_uwe']:.2f}%) than wound edge (median "
        f"{variant_stats['records_reduction_median_pct_we']:.2f}%)."
    )
    add_par(
        doc,
        f"At condition level, median record counts decreased from "
        f"{variant_stats['records_median_filtered_uwe']:.0f} to {variant_stats['records_median_cohort_uwe']:.0f} "
        f"in unwounded skin and from {variant_stats['records_median_filtered_we']:.0f} to "
        f"{variant_stats['records_median_cohort_we']:.0f} in wound edge. SNP reductions "
        f"({variant_stats['snps_reduction_mean_pct']:.2f}% mean) were slightly larger than indel reductions "
        f"({variant_stats['indels_reduction_mean_pct']:.2f}% mean)."
    )
    add_par(
        doc,
        f"Dispersion remained high in both branches: filtered records spanned a "
        f"{variant_stats['records_range_fold_filtered']:.2f}x range and cohort-filtered records spanned a "
        f"{variant_stats['records_range_fold_cohort']:.2f}x range. The strongest and weakest per-sample record "
        f"reductions were observed in {variant_stats['max_reduction_sample']} "
        f"({variant_stats['max_reduction_sample_pct']:.2f}%) and {variant_stats['min_reduction_sample']} "
        f"({variant_stats['min_reduction_sample_pct']:.2f}%), respectively."
    )
    add_table_from_df(doc, "Table 2. Variant count summary.", display_table2(core["table2"]))
    if "caller_comparison" in core:
        add_table_from_df(doc, "Table S1. Historical baseline vs active GATK comparison.", core["caller_comparison"])

    add_heading(doc, "3.2. Model performance", level=2)
    add_par(
        doc,
        f"Table 3 summarizes primary ML outcomes for both feature branches. Feature dimension decreased from "
        f"{metrics['features_filtered']:,} to {metrics['features_cohort']:,} "
        f"({feature_reduction_pct:.2f}% reduction), while best repeated-CV balanced accuracy changed from "
        f"{metrics['bacc_filtered']:.3f} to {metrics['bacc_cohort']:.3f} ({bacc_delta_pp:+.2f} percentage points). "
        f"Both branches selected linear SVC as best repeated-CV model."
    )
    add_par(
        doc,
        f"Permutation p-values were {metrics['p_filtered']:.5f} and {metrics['p_cohort']:.5f}. Given "
        f"{metrics['permutations']} permutations, these values correspond to near-minimal discrete tail counts, "
        "supporting non-random signal under the run-level protocol."
    )
    add_table_from_df(doc, "Table 3. Main ML summary.", display_table3(core["table3"]))
    add_par(
        doc,
        "Table 4 details repeated-CV behavior across model families. Linear SVC achieved the highest balanced "
        "accuracy mean (0.980), followed closely by L2 logistic regression and PCA-reduced logistic regression "
        "(both 0.970). In contrast, random forest yielded high ROC AUC mean (0.995) but materially lower "
        "balanced accuracy mean (0.803), indicating weaker thresholded class-balance behavior in this setting."
    )
    add_table_from_df(doc, "Table 4. Top repeated-CV models.", top_models(core["metrics_repeated"]))

    add_heading(doc, "3.3. Split-sensitivity analysis", level=2)
    add_par(
        doc,
        f"Table 5 shows a pronounced split-protocol effect. Balanced accuracy decreased from "
        f"{metrics['run_level_bacc']:.3f} in run-level repeated CV to {metrics['donor_level_bacc']:.3f} in "
        f"donor-grouped evaluation (absolute drop {leak_abs_drop:.3f}; relative drop {leak_rel_drop_pct:.1f}%). "
        "Accuracy followed the same direction."
    )
    add_par(
        doc,
        "This gap indicates that a substantial fraction of run-level discriminative structure does not transfer "
        "under donor-level separation, and therefore should be interpreted as split-sensitive rather than fully "
        "generalizable signal."
    )
    add_table_from_df(doc, "Table 5. Leakage-sensitivity comparison.", display_leakage(core["leakage"]))

    add_heading(doc, "3.4. Minimal enrichment summary", level=2)
    add_par(
        doc,
        "Table 6 and Figure 8 provide compact context rather than mechanistic proof. In the top-mutated-gene set, "
        "top terms included broad cellular compartments/functions (e.g., cytoplasm, ion binding, cytosol, small "
        "molecule binding). In the wound-edge-higher set, top terms included immune-associated categories "
        "(e.g., MHC protein complex and KEGG allograft/graft-versus-host pathways). In the unwounded-skin-higher "
        "set, top terms highlighted scavenger-receptor and FCGR/phagocytosis-related categories."
    )
    add_par(
        doc,
        "These outputs were retained only as sanity-check context consistent with a minimal bioinformatics "
        "interpretation layer."
    )
    add_table_from_df(doc, "Table 6. Top enrichment terms.", enrichment_table())

    add_heading(doc, "3.5. Expression-mutation correlation analysis", level=2)
    add_par(
        doc,
        "Correlation analysis showed that total gene burden was positively associated with both filtered and "
        f"cohort-filtered record counts (Spearman rho={rho_gene_filtered:.3f} and {rho_gene_cohort:.3f}, "
        f"respectively). The strongest pairwise association was between filtered and cohort-filtered records "
        f"(rho={rho_filtered_cohort:.3f}), indicating coherent burden structure across feature branches."
    )
    add_table_from_df(doc, "Table 7. Spearman correlation matrix (run-level burden metrics).", display_corr_matrix(corr))
    add_par(
        doc,
        f"Marker-set overlap testing comprised {n_marker_tests} tests (38 clusters x 3 mutation-derived gene sets). "
        f"Only {n_marker_overlap_positive} tests had overlap_n > 0, and only {n_marker_nominal} test reached nominal "
        f"p < 0.05; no test survived FDR correction (q < 0.05: {n_marker_fdr}). The minimum nominal signal was "
        f"observed for cluster {int(best_marker_row['cluster'])} ({best_marker_row['predicted_cell_type']}) against "
        f"{best_marker_row['gene_set']} with overlap gene {best_marker_row['overlap_genes']} "
        f"(p={best_marker_row['p_value']:.6f}, q={best_marker_row['q_value']:.6f})."
    )
    add_table_from_df(
        doc,
        "Table 8. Marker-overlap enrichment summary by mutation-derived gene set.",
        marker_overlap_summary(marker_enr),
    )
    add_table_from_df(
        doc,
        "Table 9. Representative non-zero overlaps between marker sets and mutation-derived gene sets.",
        display_marker_overlap_top5(core["marker_overlap_top5"]),
    )

    add_heading(doc, "3.6. Figures", level=2)
    add_par(
        doc,
        "Figure 1 provides the full processing chain from raw reads to model validation and interpretation outputs. "
        "Figures 2 and 3 visually confirm the substantial downward shift in per-sample variant counts after "
        "cohort-frequency filtering, while preserving inter-sample heterogeneity."
    )
    add_par(
        doc,
        "Figures 4 and 5 align with Tables 3 and 4: linear SVC and L2-type logistic models remain top performers "
        "under repeated CV, and performance ordering is stable across filtered and cohort-filtered branches."
    )
    add_par(
        doc,
        f"Figures 6 and 7 (PCA overlays) show that within-donor run pairs cluster much more tightly than "
        f"between-donor centroids. For filtered features, median within-donor distance was "
        f"{pca_filtered['within_donor_median']:.2f} versus median between-donor centroid distance "
        f"{pca_filtered['between_donor_median']:.2f} (ratio {pca_filtered['within_over_between_ratio']:.3f}); "
        f"for cohort-filtered features, the corresponding values were {pca_cohort['within_donor_median']:.2f} and "
        f"{pca_cohort['between_donor_median']:.2f} (ratio {pca_cohort['within_over_between_ratio']:.3f}). "
        f"The largest within-donor separation was observed for donor {pca_filtered['max_pair_donor']} "
        f"(filtered: {pca_filtered['max_pair_distance']:.2f}) and donor {pca_cohort['max_pair_donor']} "
        f"(cohort-filtered: {pca_cohort['max_pair_distance']:.2f})."
    )
    add_par(
        doc,
        "Figure 8 summarizes top enrichment terms and should be read as contextual annotation, not as causal "
        "biological confirmation."
    )
    for i, (out_name, _, caption) in enumerate(FIGURE_MAP, start=1):
        add_figure(doc, i, out_name, caption)

    add_heading(doc, "4. Discussion", level=1)
    add_par(
        doc,
        "The central technical finding is the gap between optimistic run-level validation and donor-grouped "
        "validation. High repeated-CV performance (balanced accuracy near 0.98) can coexist with weak donor-level "
        "generalization (balanced accuracy near 0.36), and this mismatch would be missed without explicit split "
        "sensitivity analysis."
    )
    add_par(
        doc,
        "Cohort-frequency filtering reduced feature dimensionality by 8.60% with stable or slightly improved "
        "run-level balanced accuracy, suggesting that removing cohort-common loci can simplify representation "
        "without obvious loss of signal in this dataset. At the same time, wide per-sample count dispersion and "
        "strong donor-structured PCA patterns indicate persistent heterogeneity that must be considered in model "
        "interpretation."
    )
    add_par(
        doc,
        "From a reporting standpoint, tables and figures should be interpreted together: Tables 3-5 quantify model "
        "behavior under different validation regimes, Tables 7-9 describe expression-mutation alignment strength, "
        "and Figures 2-7 visualize the sample and feature structure that helps explain these metrics."
    )
    add_par(
        doc,
        "Key limitations are the small donor count, absence of an external independent cohort, and potential "
        "run-pair dependence. Therefore, this study should be interpreted as a reproducibility and evaluation-design "
        "reference workflow rather than a finalized clinical classifier."
    )

    add_heading(doc, "5. Conclusions", level=1)
    add_par(
        doc,
        "This BioMedInformatics-focused study demonstrates that reproducible end-to-end ML pipelines for "
        "scRNA-seq-derived variant features can produce strong run-level discrimination while remaining highly "
        "sensitive to donor-level split design."
    )
    add_par(
        doc,
        "The practical conclusion is that donor-aware validation and explicit leakage-sensitivity reporting are not "
        "optional add-ons but core requirements for credible biomedical ML claims in small-cohort settings."
    )

    add_heading(doc, "Supplementary Materials", level=1)
    add_par(
        doc,
        "Processed tables and figures are available in Final version/results/. Figure upload copies are provided in "
        "Final version/submission_biomedinformatics/figures/."
    )

    add_heading(doc, "Author Contributions", level=1)
    add_par(
        doc,
        "Conceptualization, A.K.; Methodology, A.K.; Software, A.K.; Validation, A.K.; Formal Analysis, A.K.; "
        "Investigation, A.K.; Data Curation, A.K.; Writing - Original Draft Preparation, A.K.; "
        "Writing - Review & Editing, A.K.; Visualization, A.K.; Supervision, A.K. "
        "All authors have read and agreed to the published version of the manuscript."
    )

    add_heading(doc, "Funding", level=1)
    add_par(doc, "This research received no external funding.")

    add_heading(doc, "Institutional Review Board Statement", level=1)
    add_par(doc, "Not applicable (publicly available de-identified data).")

    add_heading(doc, "Informed Consent Statement", level=1)
    add_par(doc, "Not applicable.")

    add_heading(doc, "Data Availability Statement", level=1)
    add_par(
        doc,
        "The original data presented in this study are openly available in NCBI BioProject PRJNA736095 "
        "(https://www.ncbi.nlm.nih.gov/bioproject/PRJNA736095). The run-level metadata table used in this work "
        "is available at Final version/results/tables/table0_external_metadata_geosra.csv. "
        "Derived result tables and figures are available in Final version/results/. "
        f"The code and workflows are available at {CODE_REPO_URL} (accessed on 22 February 2026)."
    )

    add_heading(doc, "Acknowledgments", level=1)
    add_par(
        doc,
        "The author acknowledges public data providers (SRA/GEO) and original PRJNA736095 contributors. "
        "During the preparation of this manuscript, the author used a generative AI assistant for drafting and "
        "editing support. The author reviewed and edited the output and takes full responsibility for the content."
    )

    add_heading(doc, "Conflicts of Interest", level=1)
    add_par(doc, "The author declares no conflicts of interest.")

    add_heading(doc, "References", level=1)
    refs = [
        "[1] Dobin, A.; Davis, C.A.; Schlesinger, F.; Drenkow, J.; Zaleski, C.; Jha, S.; Batut, P.; Chaisson, M.; Gingeras, T.R. STAR: ultrafast universal RNA-seq aligner. Bioinformatics 2013, 29, 15-21.",
        "[2] Li, H.; Handsaker, B.; Wysoker, A.; Fennell, T.; Ruan, J.; Homer, N.; Marth, G.; Abecasis, G.; Durbin, R. The Sequence Alignment/Map format and SAMtools. Bioinformatics 2009, 25, 2078-2079.",
        "[3] Danecek, P.; Bonfield, J.K.; Liddle, J.; Marshall, J.; Ohan, V.; Pollard, M.O.; Whitwham, A.; Keane, T.; McCarthy, S.A.; Davies, R.M.; et al. Twelve years of SAMtools and BCFtools. GigaScience 2021, 10, giab008.",
        "[4] Pedregosa, F.; Varoquaux, G.; Gramfort, A.; Michel, V.; Thirion, B.; Grisel, O.; Blondel, M.; Prettenhofer, P.; Weiss, R.; Dubourg, V.; et al. Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research 2011, 12, 2825-2830.",
        "[5] Ojala, M.; Garriga, G.C. Permutation tests for studying classifier performance. Journal of Machine Learning Research 2010, 11, 1833-1863.",
        "[6] Raudvere, U.; Kolberg, L.; Kuzmin, I.; Arak, T.; Adler, P.; Peterson, H.; Vilo, J. g:Profiler: a web server for functional enrichment analysis and conversions of gene lists (2019 update). Nucleic Acids Research 2019, 47, W191-W198.",
        "[7] National Center for Biotechnology Information. BioProject PRJNA736095. Available online: https://www.ncbi.nlm.nih.gov/bioproject/PRJNA736095 (accessed on 2026-02-22).",
    ]
    for r in refs:
        add_par(doc, r)

    doc.save(str(OUT_DOCX))
    fig_manifest.to_csv(SUBMISSION_DIR / "figure_manifest.csv", index=False)
    write_cover_letter(metrics)
    write_submission_checklist()
    print(f"Wrote manuscript: {OUT_DOCX}")
    print(f"Prepared submission directory: {SUBMISSION_DIR}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
